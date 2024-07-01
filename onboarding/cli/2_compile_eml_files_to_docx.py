import docx
import html2text
import io
import os
import re

from copy import deepcopy
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from email.parser import BytesParser
from email.policy import default
from tqdm import tqdm

def extract_email_and_sequence_from_filename(filename):
    email_match = re.search(r'[\w\.-]+@[\w\.-]+', filename)
    sequence_match = re.search(r'(\d+)_Onboarding - (\d+)\.', filename)
    
    email = email_match.group(0) if email_match else None
    identifier, sequence_number = (int(sequence_match.group(1)), int(sequence_match.group(2))) if sequence_match else (None, None)
    
    return email, identifier, sequence_number

def apply_markdown_styles(doc, text_content):
    placeholder = "|||"  # Placeholder for internal use
    # Create a pattern to match bold titles, possibly spanning multiple lines
    multiline_title_pattern = re.compile(r'\*\*(?:(?!\*\*).)+\*\*', re.DOTALL)

    # Substitute line breaks within bold titles with a placeholder
    text_content = multiline_title_pattern.sub(lambda m: m.group().replace('\n', placeholder), text_content)

    # Split the text content into lines
    lines = text_content.split('\n')

    # Process each line
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == '---':  # Ignore markdown horizontal line
            i += 1
            continue

        # Detect markdown titles, now considering placeholders for line breaks
        title_pattern = re.compile(r'\*\*(.*?)\*\*')
        titles = title_pattern.findall(line)

        for title in titles:
            title_text = title.replace(placeholder, ' ').strip()  # Replace placeholder with space and strip whitespace
            doc.add_heading(title_text, level=1)  # Add the title as a heading
            line = line.replace(f"**{title}**", '')  # Remove the title markdown from the line

        # If the line is empty after removing titles, skip to the next line
        if not line.strip():
            i += 1
            continue

        # Handle the Google Maps link
        if line.startswith('[Map'):
            # The Google Maps URL is expected to be on the next line
            if i + 1 < len(lines):
                next_line = lines[i + 1]
                map_url_match = re.search(r'\((https?://[^)]+)\)', next_line)
                if map_url_match:
                    # Create a hyperlink for the Google Maps URL
                    url = map_url_match.group(1)
                    p = doc.add_paragraph()
                    add_hyperlink(p, url, "Map It")
                    i += 2  # Skip the next line as it's part of the Google Maps link
                    continue

        # After processing titles and Google Maps link, handle any remaining content
        process_line_for_links(doc, line.strip())
        i += 1

    # Remove any remaining placeholders from the entire content before finalizing the document
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, ' ')

def process_line_for_links(doc, line):
    # Regular expression to match markdown links and angle bracket links
    markdown_link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    angle_bracket_link_pattern = r'<([^>]+)>'

    # Combine the patterns to match either markdown link or angle bracket link
    combined_pattern = f'({markdown_link_pattern})|({angle_bracket_link_pattern})'

    # Find all matches in the line for the combined pattern
    all_links = re.findall(combined_pattern, line)

    # Create a new paragraph
    p = doc.add_paragraph()
    last_idx = 0

    for match in all_links:
        # Extract text and url for markdown link or angle bracket link
        text, url = (match[1], match[2]) if match[1] else (match[4], match[4])

        # Add text before the hyperlink
        start_idx = line.find(match[0], last_idx)
        if start_idx != -1:
            p.add_run(line[last_idx:start_idx])

        # Add the hyperlink
        add_hyperlink(p, url, text)

        # Update last index to end of the match
        last_idx = start_idx + len(match[0])

    # Add any remaining text after the last link
    remaining_text = line[last_idx:]

    # Check for a duplicate URL following the pattern "| <URL>"
    duplicate_pattern = r'\|\s*<https?://[^>]+>'
    remaining_text = re.sub(duplicate_pattern, '', remaining_text)

    # Add the remaining text if there is any
    if remaining_text.strip():
        p.add_run(remaining_text)

def convert_remaining_urls_to_hyperlinks(doc):
    url_pattern = re.compile(r'https?://[^\s]+')

    for paragraph in doc.paragraphs:
        new_runs = []
        last_idx = 0

        for match in url_pattern.finditer(paragraph.text):
            start_idx, end_idx = match.span()
            # Add text before the URL as a regular run
            new_runs.append((paragraph.text[last_idx:start_idx], False))
            # Add the URL as a hyperlink
            new_runs.append((match.group(), True))
            last_idx = end_idx

        # Add any remaining text after the last URL
        if last_idx < len(paragraph.text):
            new_runs.append((paragraph.text[last_idx:], False))

        # If URLs were found, clear the paragraph and add new runs
        if last_idx > 0:
            paragraph.clear()
            for text, is_hyperlink in new_runs:
                if is_hyperlink:
                    add_hyperlink(paragraph, text, text)
                else:
                    paragraph.add_run(text)

def add_hyperlink(paragraph, url, text, color='0000FF', underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element (a new run) that will contain the hyperlink
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')  # Create run properties

    # Set the style for the hyperlink (color, underline)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)  # Set the color
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object for the paragraph and append the hyperlink
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # Set the font size, color, and underline properties using the python-docx API
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 0, 255)  # Set the font color to blue
    if underline:
        r.font.underline = True

# Catch all the other links coming from the text in the Other Links section and anywhere else there's a links comgin fro mthe EML files
def process_other_links_section(doc, text_content):
    lines = text_content.split('\n')
    other_links_section_found = False

    for line in lines:
        if "Other Links" in line:
            other_links_section_found = True
            continue

        if other_links_section_found and line.strip():
            # Improved URL extraction using regex
            urls = re.findall(r'https?://[^\s]+', line)
            for url in urls:
                add_hyperlink(doc.add_paragraph(), url, url)
        elif other_links_section_found and not line.strip():
            break
                                
def manage_header_relationships(template_doc, target_doc):
    template_part = template_doc.sections[0].header.part
    target_part = target_doc.sections[0].header.part

    for rel in template_part.rels.values():
        # Check if the relationship is an image
        if rel.reltype == RT.IMAGE:
            # Access the image blob directly from the package part
            image_blob = rel.target_part._blob

            # Add the image to the target document and get the new relationship ID
            image_stream = io.BytesIO(image_blob)
            new_rel_id, _ = target_part.get_or_add_image(image_stream)  # Extract rel ID from the tuple

            # Replace old rel ID with new rel ID in header XML
            for element in target_doc.sections[0].header._element:
                xml_str = element.xml
                updated_xml_str = xml_str.replace(rel.rId, new_rel_id)
                new_element = parse_xml(updated_xml_str.encode())
                element.getparent().replace(element, new_element)

# Get rid of the pipe characters coming from the emails
def remove_pipe_characters(doc):
    for paragraph in doc.paragraphs:
        if '|' in paragraph.text:
            paragraph.text = paragraph.text.replace('|', '')  # Replace '|' with an empty string

# Apply the header from the branding template DOCX file
def apply_header_from_template(doc, template_path):
    template_doc = Document(template_path)
    template_header = template_doc.sections[0].header

    for section in doc.sections:
        # Clear existing header elements
        section.header._element.clear()

        # Copy header XML from template
        for element in template_header._element:
            section.header._element.append(deepcopy(element))

        # Manage relationships (images, etc.)
        manage_header_relationships(template_doc, doc)
        
def add_title_to_first_page(doc, title_text, font_size=24):
    """
    Adds a title to the first page of the document and center aligns it.

    :param doc: The Document object to which the title will be added.
    :param title_text: The text of the title.
    :param font_size: The font size of the title.
    """
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align the paragraph
    title_run = title.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(font_size)
    title.style = doc.styles['Title']  # Optionally, use the built-in 'Title' style
    doc.add_paragraph()  # Add a blank line after the title

# This function will be called with the text content you extract from the emails
                             
def compile_emails_to_docx(email_downloads_folder):
    email_files_map = {}
    email_already_written = {}  # Track if the email has already been written for each email address

    for filename in os.listdir(email_downloads_folder):
        email, identifier, sequence_number = extract_email_and_sequence_from_filename(filename)
        if email and sequence_number is not None:
            file_info = (identifier, os.path.join(email_downloads_folder, filename))
            if email not in email_files_map:
                email_files_map[email] = {}
                first_instance_email_line = None  # Ensure the line storage is reset for each email
            if sequence_number not in email_files_map[email] or email_files_map[email][sequence_number][0] < identifier:
                email_files_map[email][sequence_number] = file_info

    for email, files in tqdm(email_files_map.items(), desc="Processing emails"):
        doc = Document()
        apply_header_from_template(doc, './header_long_story_short.docx')

        # Add the title to the first page
        add_title_to_first_page(doc, "Onboarding Questionnaire Results", font_size=24)

        sorted_files = sorted(files.values(), key=lambda x: (x[1], x[0]))
        email_already_written[email] = False

        for _, file in tqdm(sorted_files, desc=f"Compiling for {email}", leave=False):
            with open(file, 'rb') as f:
                email_msg = BytesParser(policy=default).parse(f)
                html_part = email_msg.get_body(preferencelist=('html'))
                
                if html_part is not None:
                    html_content = html_part.get_content()
                    # Convert the HTML content to text with html2text
                    text_content = html2text.html2text(html_content)

                    # Check if the email line has already been written for this email address
                    if email_already_written[email]:
                        # If so, remove the last occurrence of the email line
                        # which should be the last line in the converted text
                        lines = text_content.strip().split('\n')
                        if 'Email' in lines[-2]:  # Assuming 'Email' is the second last line
                            # Remove the last two lines which contain the 'Email' header and email address
                            text_content = '\n'.join(lines[:-2])
                    else:
                        # Mark the email as written
                        email_already_written[email] = True

                    # Apply markdown styles
                    apply_markdown_styles(doc, text_content)

                    # Process other links
                    process_other_links_section(doc, text_content)

                    # Convert remaining URLs to hyperlinks
                    convert_remaining_urls_to_hyperlinks(doc)

                else:
                    tqdm.write(f"No HTML content in {file}")

        # Remove pipe characters
        remove_pipe_characters(doc)

        # Assuming 'email' is a string containing the email filename with the '.eml' extension
        email_without_extension = email.replace('.eml', '')

        # Now using the modified email string to form the document filename
        doc_filename = f"docs/{email_without_extension}.docx"
        doc.save(doc_filename)
        tqdm.write(f"Saved {doc_filename}")

# Folder settings
email_downloads_folder = 'email_downloads'
compile_emails_to_docx(email_downloads_folder)