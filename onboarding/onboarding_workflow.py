import streamlit as st
import imaplib
import email
import os
from email.header import decode_header
import json
import docx
import html2text
import io
import re
from copy import deepcopy
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn
from email.parser import BytesParser
from email.policy import default
import requests
import PyPDF2

# Load configuration
@st.cache_resource
def load_config():
    config_path = os.path.join(os.path.dirname(__file__), 'config.json')
    if os.path.exists(config_path):
        with open(config_path, 'r') as config_file:
            return json.load(config_file)
    else:
        return {
            "IMAP_SERVER": "",
            "EMAIL_ADDRESS": "",
            "PASSWORD": ""
        }

def save_config(imap_server, email_address, password):
    config = {
        "IMAP_SERVER": imap_server,
        "EMAIL_ADDRESS": email_address,
        "PASSWORD": password
    }
    config_path = os.path.join(os.path.dirname(__file__), 'config.json')
    with open(config_path, 'w') as config_file:
        json.dump(config, config_file)

config = load_config()
IMAP_SERVER = config['IMAP_SERVER']
EMAIL_ADDRESS = config['EMAIL_ADDRESS']
PASSWORD = config['PASSWORD']
LOCAL_FOLDER = os.path.join(os.path.dirname(__file__), 'email_downloads')
DOCS_FOLDER = os.path.join(os.path.dirname(__file__), 'docs')
HEADER_TEMPLATE = os.path.join(os.path.dirname(__file__), 'header_long_story_short.docx')

# Ollama Setup
model = 'llama3:8b-instruct-fp16'
ollama_url = 'http://localhost:11434/api/generate'

# Streamlit app
st.title("Onboarding GAP Analysis Generator")

# Workflow Execution
def run_workflow():
    status_placeholder = st.empty()
    progress_bar = st.progress(0)
    
    def update_status(message, progress=None):
        if progress is not None:
            progress_bar.progress(progress)
        status_placeholder.text(message)
    
    def check_and_download_emails():
        mail = imaplib.IMAP4_SSL(IMAP_SERVER)
        mail.login(EMAIL_ADDRESS, PASSWORD)
        mail.select('INBOX')
        status, email_ids = mail.search(None, 'UNSEEN')
        
        if status == 'OK':
            email_id_list = email_ids[0].split()
            for i, email_id in enumerate(email_id_list):
                status, email_data = mail.fetch(email_id, '(RFC822)')
                if status == 'OK':
                    raw_email = email_data[0][1]
                    msg = email.message_from_bytes(raw_email)
                    subject, encoding = decode_header(msg["Subject"])[0]
                    if isinstance(subject, bytes):
                        subject = subject.decode(encoding or 'utf-8')
                    filename = f"{email_id.decode('utf-8')}_{subject}.eml"
                    local_path = os.path.join(LOCAL_FOLDER, filename)
                    os.makedirs(os.path.dirname(local_path), exist_ok=True)
                    
                    if not os.path.exists(local_path):
                        with open(local_path, 'wb') as email_file:
                            email_file.write(raw_email)
                        update_status(f"Downloaded: {filename}", (i + 1) / len(email_id_list))
                    else:
                        update_status(f"Skipped: {filename} (already exists)", (i + 1) / len(email_id_list))
        
        mail.close()
        mail.logout()
        st.success("Email check and download complete!")

    def extract_email_and_sequence_from_filename(filename):
        email_match = re.search(r'[\w\.-]+@[\w\.-]+', filename)
        sequence_match = re.search(r'(\d+)_Onboarding - (\d+)\.', filename)
        email = email_match.group(0) if email_match else None
        identifier, sequence_number = (int(sequence_match.group(1)), int(sequence_match.group(2))) if sequence_match else (None, None)
        return email, identifier, sequence_number

    def apply_markdown_styles(doc, text_content):
        placeholder = "|||"
        multiline_title_pattern = re.compile(r'\*\*(?:(?!\*\*).)+\*\*', re.DOTALL)
        text_content = multiline_title_pattern.sub(lambda m: m.group().replace('\n', placeholder), text_content)
        lines = text_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            if line.strip() == '---':
                i += 1
                continue
            title_pattern = re.compile(r'\*\*(.*?)\*\*')
            titles = title_pattern.findall(line)
            for title in titles:
                title_text = title.replace(placeholder, ' ').strip()
                doc.add_heading(title_text, level=1)
                line = line.replace(f"**{title}**", '')
            if not line.strip():
                i += 1
                continue
            if line.startswith('[Map'):
                if i + 1 < len(lines):
                    next_line = lines[i + 1]
                    map_url_match = re.search(r'\((https?://[^)]+)\)', next_line)
                    if map_url_match:
                        url = map_url_match.group(1)
                        p = doc.add_paragraph()
                        add_hyperlink(p, url, "Map It")
                        i += 2
                        continue
            process_line_for_links(doc, line.strip())
            i += 1
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, ' ')

    def process_line_for_links(doc, line):
        markdown_link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
        angle_bracket_link_pattern = r'<([^>]+)>'
        combined_pattern = f'({markdown_link_pattern})|({angle_bracket_link_pattern})'
        all_links = re.findall(combined_pattern, line)
        p = doc.add_paragraph()
        last_idx = 0
        for match in all_links:
            text, url = (match[1], match[2]) if match[1] else (match[4], match[4])
            start_idx = line.find(match[0], last_idx)
            if start_idx != -1:
                p.add_run(line[last_idx:start_idx])
            add_hyperlink(p, url, text)
            last_idx = start_idx + len(match[0])
        remaining_text = line[last_idx:]
        duplicate_pattern = r'\|\s*<https?://[^>]+>'
        remaining_text = re.sub(duplicate_pattern, '', remaining_text)
        if remaining_text.strip():
            p.add_run(remaining_text)

    def add_hyperlink(paragraph, url, text, color='0000FF', underline=True):
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)
        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        r = paragraph.add_run()
        r._r.append(hyperlink)
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 255)
        if underline:
            r.font.underline = True

    def process_other_links_section(doc, text_content):
        lines = text_content.split('\n')
        other_links_section_found = False
        for line in lines:
            if "Other Links" in line:
                other_links_section_found = True
                continue
            if other_links_section_found and line.strip():
                urls = re.findall(r'https?://[^\s]+', line)
                for url in urls:
                    add_hyperlink(doc.add_paragraph(), url, url)
            elif other_links_section_found and not line.strip():
                break

    def manage_header_relationships(template_doc, target_doc):
        template_part = template_doc.sections[0].header.part
        target_part = target_doc.sections[0].header.part
        for rel in template_part.rels.values():
            if rel.reltype == RT.IMAGE:
                image_blob = rel.target_part._blob
                image_stream = io.BytesIO(image_blob)
                new_rel_id, _ = target_part.get_or_add_image(image_stream)
                for element in target_doc.sections[0].header._element:
                    xml_str = element.xml
                    updated_xml_str = xml_str.replace(rel.rId, new_rel_id)
                    new_element = parse_xml(updated_xml_str.encode())
                    element.getparent().replace(element, new_element)

    def remove_pipe_characters(doc):
        for paragraph in doc.paragraphs:
            if '|' in paragraph.text:
                paragraph.text = paragraph.text.replace('|', '')

    def apply_header_from_template(doc, template_path):
        template_doc = Document(template_path)
        template_header = template_doc.sections[0].header
        for section in doc.sections:
            section.header._element.clear()
            for element in template_header._element:
                section.header._element.append(deepcopy(element))
            manage_header_relationships(template_doc, doc)

    def add_title_to_first_page(doc, title_text, font_size=24):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(title_text)
        title_run.bold = True
        title_run.font.size = Pt(font_size)
        title.style = doc.styles['Title']
        doc.add_paragraph()

    def compile_emails_to_docx():
        email_files_map = {}
        email_already_written = {}
        
        for filename in os.listdir(LOCAL_FOLDER):
            email, identifier, sequence_number = extract_email_and_sequence_from_filename(filename)
            if email and sequence_number is not None:
                file_info = (identifier, os.path.join(LOCAL_FOLDER, filename))
                if email not in email_files_map:
                    email_files_map[email] = {}
                if sequence_number not in email_files_map[email] or email_files_map[email][sequence_number][0] < identifier:
                    email_files_map[email][sequence_number] = file_info
        
        for i, (email, files) in enumerate(email_files_map.items()):
            doc = Document()
            apply_header_from_template(doc, HEADER_TEMPLATE)
            add_title_to_first_page(doc, "Onboarding Questionnaire Results", font_size=24)
            
            sorted_files = sorted(files.values(), key=lambda x: (x[1], x[0]))
            email_already_written[email] = False
            
            for _, file in sorted_files:
                with open(file, 'rb') as f:
                    email_msg = BytesParser(policy=default).parse(f)
                    html_part = email_msg.get_body(preferencelist=('html'))
                    
                    if html_part is not None:
                        html_content = html_part.get_content()
                        text_content = html2text.html2text(html_content)
                        
                        if email_already_written[email]:
                            lines = text_content.strip().split('\n')
                            if 'Email' in lines[-2]:
                                text_content = '\n'.join(lines[:-2])
                        else:
                            email_already_written[email] = True
                        
                        apply_markdown_styles(doc, text_content)
                        process_other_links_section(doc, text_content)
                    else:
                        update_status(f"No HTML content in {file}")
            
            remove_pipe_characters(doc)
            email_without_extension = email.replace('.eml', '')
            doc_filename = f"{DOCS_FOLDER}/{email_without_extension}.docx"
            doc.save(doc_filename)
            update_status(f"Saved {doc_filename}", (i + 1) / len(email_files_map))
        
        st.success("Email compilation complete!")

    def read_word_document(file_path):
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])

    def read_pdf_document(file_path):
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            return '\n'.join([page.extract_text() for page in reader.pages])

    def analyze_with_ollama(content, prompt_template, model=model, max_tokens=300):
        prompt = prompt_template.replace("{content}", content)
        
        def generate_response(prompt, context=[]):
            data = {
                'model': model,
                'prompt': prompt,
                'context': context,
                'max_tokens': max_tokens
            }
            response = requests.post(ollama_url, json=data, stream=True)
            response.raise_for_status()
            accumulated_response = ""
            for line in response.iter_lines():
                if line:
                    body = json.loads(line)
                    response_part = body.get('response', '')
                    accumulated_response += response_part
                    if len(accumulated_response) > 100:  # Yield only when accumulated response is sufficiently large
                        yield accumulated_response
                        accumulated_response = ""
                    if body.get('done', False):
                        yield accumulated_response  # Yield any remaining part when done
                        break
                    if 'error' in body:
                        raise Exception(body['error'])
        
        return generate_response(prompt)

    def write_analysis_to_file(analysis, file_path):
        with open(file_path, 'w') as file:
            file.write(analysis)

    def create_gap_reports():
        prompt_template = "Analyze the following text and provide a GAP analysis report. Also give me a summary of missing information with possible solutions based on our marketing services. Order the possible solutions you suggest at the end in the order they are best executed in. In the GAP analysis report, try to guide them with judgement based on what they entered even if the information is insufficient. Be sure to provide feedback for all nine sections of the marketing plan process: {content}"
        
        files = [f for f in os.listdir(DOCS_FOLDER) if f.endswith('.docx') or f.endswith('.pdf')]
        
        for i, file in enumerate(files):
            file_path = os.path.join(DOCS_FOLDER, file)
            update_status(f"Processing {file_path}...", (i + 1) / len(files))
            
            if file.endswith('.docx'):
                content = read_word_document(file_path)
            elif file.endswith('.pdf'):
                content = read_pdf_document(file_path)
            
            analysis_generator = analyze_with_ollama(content, prompt_template)
            
            analysis_text = ""
            st.write_stream(analysis_generator)
            for part in analysis_generator:
                analysis_text = part
            
            analysis_file_path = os.path.join(DOCS_FOLDER, f"{os.path.splitext(file)[0]}_GAP_analysis.txt")
            write_analysis_to_file(analysis_text, analysis_file_path)
            
            update_status(f"Analysis report created: {analysis_file_path}", (i + 1) / len(files))
        
        st.success("GAP report generation complete!")

    # Run all steps
    check_and_download_emails()
    compile_emails_to_docx()
    create_gap_reports()

def run_onboarding_workflow():
    # Sidebar for configuration
    with st.sidebar.expander("API Configuration", expanded=False):
        IMAP_SERVER = st.text_input("Enter IMAP Server", value=config.get("IMAP_SERVER", ""))
        EMAIL_ADDRESS = st.text_input("Enter Email Address", value=config.get("EMAIL_ADDRESS", ""))
        PASSWORD = st.text_input("Enter Password", value=config.get("PASSWORD", ""), type="password")

        if st.button("Save API Settings"):
            save_config(IMAP_SERVER, EMAIL_ADDRESS, PASSWORD)
            st.success("API settings saved!")

    if st.button("Run Workflow"):
        run_workflow()

    # Display Results
    st.header("Results")
    if os.path.exists(DOCS_FOLDER):
        st.subheader("Generated Documents:")
        for file in os.listdir(DOCS_FOLDER):
            if file.endswith('.docx') or file.endswith('.txt'):
                file_path = os.path.join(DOCS_FOLDER, file)
                with open(file_path, "rb") as f:
                    st.download_button(
                        label=f"Download {file}",
                        data=f,
                        file_name=file
                    )
    else:
        st.write("No documents generated yet.")

    # Main execution
    if not os.path.exists(LOCAL_FOLDER):
        os.makedirs(LOCAL_FOLDER)
    if not os.path.exists(DOCS_FOLDER):
        os.makedirs(DOCS_FOLDER)
